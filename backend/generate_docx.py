"""
DOCX generator — builds Word documents from schema + form data.
Uses python-docx with IEDC branding (black + red theme).
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


RED = RGBColor(0xDC, 0x26, 0x26)
BLACK = RGBColor(0x11, 0x11, 0x11)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_RED = RGBColor(0xFE, 0xF2, 0xF2)


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_header(doc: Document):
    p = doc.add_paragraph()
    run = p.add_run("INNOVATION & ENTREPRENEURSHIP DEVELOPMENT CELL")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RED
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p2 = doc.add_paragraph()
    r2 = p2.add_run("Affiliated to Kerala Startup Mission (KSUM)  |  Government of Kerala")
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Horizontal rule via border
    pPr = p2._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "double")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "DC2626")
    pBdr.append(bottom)
    pPr.append(pBdr)

    doc.add_paragraph()


def _add_title(doc: Document, title: str):
    p = doc.add_paragraph()
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RED
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def _add_field(doc: Document, label: str, value: str):
    p = doc.add_paragraph()
    label_run = p.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.size = Pt(11)
    val_run = p.add_run(str(value) if value else "—")
    val_run.font.size = Pt(11)


def _add_table_field(doc: Document, field: dict, rows_data: list):
    columns = field.get("columns", [])
    if not columns:
        return

    label_p = doc.add_paragraph()
    lr = label_p.add_run(field["label"])
    lr.bold = True
    lr.font.size = Pt(11)

    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0]
    for i, col in enumerate(columns):
        cell = hdr.cells[i]
        cell.text = col.get("label", col["name"])
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        _set_cell_bg(cell, "DC2626")

    # Data rows
    for idx, row_data in enumerate(rows_data):
        row = table.add_row()
        for i, col in enumerate(columns):
            col_name = col["name"]
            val = row_data.get(col_name, "")
            # Compute total for budget tables
            if col.get("computed") and col_name == "total":
                try:
                    qty = float(row_data.get("quantity", 0) or 0)
                    uc = float(row_data.get("unit_cost", 0) or 0)
                    val = f"{qty * uc:.2f}"
                except (ValueError, TypeError):
                    val = "0.00"
            cell = row.cells[i]
            cell.text = str(val) if val else "—"
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            if idx % 2 == 1:
                _set_cell_bg(cell, "FEF2F2")

    doc.add_paragraph()


def _add_signatures(doc: Document, left_name: str, left_title: str,
                    right_name: str = "IEDC Nodal Officer",
                    right_title: str = "Innovation & Entrepreneurship Development Cell"):
    doc.add_paragraph()
    doc.add_paragraph()
    table = doc.add_table(rows=2, cols=2)

    def _sig_cell(cell, name, title):
        p = cell.paragraphs[0]
        run = p.add_run("_" * 28)
        run.font.size = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2 = cell.add_paragraph(name)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.runs[0].bold = True
        p3 = cell.add_paragraph(title)
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.runs[0].font.size = Pt(9)
        p3.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    _sig_cell(table.rows[0].cells[0], left_name, left_title)
    _sig_cell(table.rows[0].cells[1], right_name, right_title)


def build_docx(schema: dict, data: dict, output_path: str):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    _add_header(doc)
    _add_title(doc, schema.get("name", "Document"))

    coordinator = data.get("coordinator", "—")

    for field in schema.get("fields", []):
        field_name = field["name"]
        field_type = field.get("type", "text")
        value = data.get(field_name)

        if field_type == "table":
            rows_data = value if isinstance(value, list) else []
            _add_table_field(doc, field, rows_data)
        else:
            if value:
                _add_field(doc, field["label"], str(value))

    _add_signatures(
        doc,
        left_name=coordinator,
        left_title="Event Coordinator",
    )

    doc.save(output_path)
